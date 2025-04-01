# Export/export_word.py
from docx import Document
import streamlit as st
import tempfile

def generate_procurement_word(results: list) -> None:
    """
    Exporterar en sammanställning av analyserade data till ett Word-dokument.
    
    Args:
        results (list): Lista med dictionaries innehållande 'filename' och 'data'.
    """
    doc = Document()
    doc.add_heading("Jämförelse av Försäkringsdokument", 0)
    
    for r in results:
        data = r.get("data", {})
        doc.add_heading(r.get("filename", "Utan filnamn"), level=1)
    
        table = doc.add_table(rows=1, cols=2)
        table.style = "Light List Accent 1"
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Fält"
        hdr_cells[1].text = "Värde"
    
        for label, key in [
            ("Poäng", "score"),
            ("Premie", "premie"),
            ("Självrisk", "självrisk"),
            ("Maskiner", "maskiner"),
            ("Varor", "varor"),
            ("Byggnad", "byggnad"),
            ("Transport", "transport"),
            ("Produktansvar", "produktansvar"),
            ("Ansvar", "ansvar"),
            ("Rättsskydd", "rättsskydd"),
            ("GDPR ansvar", "gdpr_ansvar"),
            ("Karens", "karens"),
            ("Ansvarstid", "ansvarstid")
        ]:
            row_cells = table.add_row().cells
            row_cells[0].text = label
            row_cells[1].text = str(data.get(key, r.get(key, "saknas")))
    
        doc.add_paragraph("")
    
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        doc.save(tmp.name)
        with open(tmp.name, "rb") as file:
            st.download_button("📄 Ladda ner Word", data=file.read(), file_name="forsakringsjämforelse.docx")
